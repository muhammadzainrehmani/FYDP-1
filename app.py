import streamlit as st
import os
from PyPDF2 import PdfReader
import docx
from langchain.vectorstores import Qdrant
import random
from datetime import datetime
from langchain import PromptTemplate
from langchain.chains import RetrievalQA
import string
from qdrant_client import QdrantClient
from dotenv import load_dotenv
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from streamlit_chat import message
from langchain.callbacks import get_openai_callback
from langchain.docstore.document import Document
from langchain_google_genai import ChatGoogleGenerativeAI
import google.generativeai as genai
from google.api_core import retry

# Store secrets securely
genai_api_key = st.secrets["genai_api_key"]

# Configure the Generative AI API with the API key
genai.configure(api_key=genai_api_key)

qdrant_url = st.secrets["QDRANT_URL"]
qdrant_api_key = st.secrets["QDRANT_API_KEY"]

# google_creds_path = st.secrets["google_application_credentials"]
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")

# When running lots of queries, it's a good practice to use a retry policy so your code
# automatically retries when hitting Resource Exhausted (quota limit) error
retry_policy = {
    "retry": retry.Retry(predicate=retry.if_transient_error, initial=10, multiplier=1.5, timeout=300)
}

# "with" notation
def main():
    load_dotenv()
    st.set_page_config(page_title="General Physician Assistant")
    st.header("GP-Assist")

    if "conversation" not in st.session_state:
        st.session_state.conversation = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "processComplete" not in st.session_state:
        st.session_state.processComplete = None

    with st.sidebar:
        uploaded_files =  st.file_uploader("Upload your Medical Files",type=['pdf','Docx','Radioactive Image'],accept_multiple_files=True)
        openai_api_key = genai_api_key
        # openai_api_key = st.text_input("OpenAI API Key", key=openapi_key , type="password")
        process = st.button("Process")
    if process:
        if not openai_api_key:
            st.info("Please add your OpenAI API key to continue.")
            st.stop()

        text_chunks_list = []
        for uploaded_file in uploaded_files:
            file_name = uploaded_file.name
            file_text = get_files_text(uploaded_file)
            # get text chunks
            text_chunks = get_text_chunks(file_text, file_name )
            text_chunks_list.extend(text_chunks)
            # create vetore stores
        curr_date = str(datetime.now())
        collection_name = "".join(random.choices(string.ascii_letters, k=4)) + curr_date.split('.')[0].replace(':', '-').replace(" ", 'T')
        vectorestore = get_vectorstore(text_chunks_list, collection_name)
        st.write("Vectore Store Created...")
        # create qa chain
        num_chunks = 4
        st.session_state.conversation = get_qa_chain(vectorestore,num_chunks) #for openAI

        st.session_state.processComplete = True

    if  st.session_state.processComplete == True:
        user_question = st.chat_input("Submit your symptoms and inquire about your medical reports.")
        if user_question:
            handel_userinput(user_question)

# Function to get the input file and read the text from it.
def get_files_text(uploaded_file):
    text = ""
    split_tup = os.path.splitext(uploaded_file.name)
    file_extension = split_tup[1]
    if file_extension == ".pdf":
        text += get_pdf_text(uploaded_file)
    elif file_extension == ".docx":
        text += get_docx_text(uploaded_file)
    else:
        pass
    return text

# Function to read PDF Files
def get_pdf_text(pdf):
    pdf_reader = PdfReader(pdf)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text

def get_docx_text(file):
    doc = docx.Document(file)
    allText = []
    for docpara in doc.paragraphs:
        allText.append(docpara.text)
    text = ' '.join(allText)
    return text



def get_text_chunks(text, filename):
    # spilit ito chuncks
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=80,
        chunk_overlap=20,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    doc_list = []
    for chunk in chunks:
        metadata = {"source": filename}
        doc_string = Document(page_content=chunk, metadata=metadata)
        doc_list.append(doc_string)
    return doc_list


def get_vectorstore(text_chunks, COLLECTION_NAME):
    # Using the hugging face embedding models
    try:
        # creating the Vectore Store using Facebook AI Semantic search
        knowledge_base = Qdrant.from_documents(
            documents = text_chunks,
            embedding = embeddings,
            url=qdrant_url,
            prefer_grpc=True,
            api_key=qdrant_api_key,
            collection_name=COLLECTION_NAME,
        )
    except Exception as e:
        st.write(f"Error: {e}")
    return knowledge_base

def get_qa_chain(vectorstore,num_chunks):
    # prompt_template = """
    # You are a medical expert. Based on the provided context—which includes documents detailing doctor intelligence and previous cases handled by doctors—generate treatment recommendations for the user's symptoms. Your response should include medication suggestions along with dosage and administration instructions when applicable. If no relevant information is found in the context, return "N/A"; otherwise, provide a precise answer in 200 words.
    # Context: {context}
    # Question: {question}"""
    prompt_template = """
    You are a highly specialized medical consultant with access to a repository of previous cases and medical intelligence. Based on the patient's symptoms, generate a comprehensive diagnostic support report that includes only the following elements:
    - Disease explanations,
    - Care recommendations,
    - Tailored advice on precautions,
    - Medication suggestions with detailed dosage and administration instructions,
    - Workout recommendations, and
    - Diet recommendations.

    Do not include any extraneous commentary or reveal any underlying context or source data. If no relevant information is available from the context, generate the recommendations based solely on general best practices.
    Context: {context}
    Question: {question}"""



    mprompt_url = PromptTemplate(
        template=prompt_template, input_variables=["context", "question"], validate_template=False)
    chain_type_kwargs = {"prompt": mprompt_url}


    # qa = RetrievalQA.from_chain_type(llm=ChatOpenAI(model = "gpt-3.5-turbo"), chain_type="stuff",
    #                             retriever=vectorstore.as_retriever(search_type="similarity",
    #                                                         search_kwargs={"k": num_chunks}), chain_type_kwargs=chain_type_kwargs, return_source_documents=True)
    qa = RetrievalQA.from_chain_type(llm = ChatGoogleGenerativeAI(
    model="gemini-1.5-flash", temperature=1, max_tokens=200, api_key=genai_api_key, request_options=retry_policy), chain_type="stuff",
                                retriever=vectorstore.as_retriever(search_type="similarity",
                                                            search_kwargs={"k": num_chunks}), chain_type_kwargs=chain_type_kwargs, return_source_documents=True)
    return qa


def handel_userinput(user_question):
    with st.spinner('Generating response...'):
        result = st.session_state.conversation({"query": user_question})
        response = result['result']
        source = result['source_documents'][0].metadata['source']
    st.session_state.chat_history.append(user_question)
    st.session_state.chat_history.append(f"{response} \n Source Document: {source}")


    # Layout of input/response containers
    response_container = st.container()

    with response_container:
        for i, messages in enumerate(st.session_state.chat_history):
            if i % 2 == 0:
                message(messages, is_user=True, key=str(i))
            else:
                message(messages, key=str(i))


if __name__ == '__main__':
    main()






