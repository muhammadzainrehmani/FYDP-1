import streamlit as st
import os
from PyPDF2 import PdfReader
import docx
from langchain.vectorstores import Qdrant
from datetime import datetime
from langchain import PromptTemplate
from langchain.chains import RetrievalQA
import string
from qdrant_client import QdrantClient
from dotenv import load_dotenv
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from streamlit_chat import message
from langchain.docstore.document import Document
from langchain_google_genai import ChatGoogleGenerativeAI
import google.generativeai as genai
from google.api_core import retry

# Store secrets securely
genai_api_key = st.secrets["genai_api_key"]

# Configure the Generative AI API with the API key
genai.configure(api_key=genai_api_key)

qdrant_url = st.secrets["QDRANT_URL"]
qdrant_api_key = st.secrets["QDRANT_API_KEY"]

# Initialize the embeddings model
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")

# Setup a retry policy
retry_policy = {
    "retry": retry.Retry(predicate=retry.if_transient_error, initial=10, multiplier=1.5, timeout=300)
}

# Use a fixed collection name for persistence across sessions
COLLECTION_NAME = "gp_assist_collection"

def main():
    load_dotenv()
    st.set_page_config(page_title="General Physician Assistant")
    st.header("GP-Assist")

    # Initialize session state
    if "conversation" not in st.session_state:
        st.session_state.conversation = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "processComplete" not in st.session_state:
        st.session_state.processComplete = False

    # Sidebar: Allow file uploads
    with st.sidebar:
        uploaded_files = st.file_uploader("Upload your Medical Files", 
                                          type=['pdf', 'docx', 'Radioactive Image'],
                                          accept_multiple_files=True)
        process = st.button("Process")

    # Create a QdrantClient to check for an existing collection
    client = QdrantClient(url=qdrant_url, api_key=qdrant_api_key)
    try:
        collections = client.get_collections().collections
        existing_collection_names = [col.name for col in collections]
    except Exception as e:
        st.write(f"Error fetching collections: {e}")
        existing_collection_names = []

    # Process any uploaded documents into text chunks
    text_chunks_list = []
    if uploaded_files:
        for uploaded_file in uploaded_files:
            file_name = uploaded_file.name
            file_text = get_files_text(uploaded_file)
            text_chunks = get_text_chunks(file_text, file_name)
            text_chunks_list.extend(text_chunks)

    if process:
        if not genai_api_key:
            st.info("Please add your OpenAI API key to continue.")
            st.stop()

        # If the collection already exists, load it
        if COLLECTION_NAME in existing_collection_names:
            try:
                vectorestore = Qdrant.from_existing_collection(
                    embedding=embeddings,
                    collection_name=COLLECTION_NAME,
                    url=qdrant_url,
                    prefer_grpc=True,
                    api_key=qdrant_api_key
                )
                st.write("Existing vector store loaded.")
                # If new documents were uploaded, add them to the collection
                if text_chunks_list:
                    try:
                        vectorestore.add_documents(text_chunks_list)
                        st.write("New documents added to the existing collection.")
                    except Exception as e:
                        st.write(f"Error adding documents: {e}")
                else:
                    st.write("No new documents uploaded; using the existing collection.")
            except Exception as e:
                st.write(f"Error loading existing vector store: {e}")
                st.stop()
        else:
            # No collection exists yet. If no documents were uploaded, show an info message.
            if not text_chunks_list:
                st.info("No existing collection found. Please upload documents to create one.")
                st.stop()
            try:
                vectorestore = Qdrant.from_documents(
                    documents=text_chunks_list,
                    embedding=embeddings,
                    url=qdrant_url,
                    prefer_grpc=True,
                    api_key=qdrant_api_key,
                    collection_name=COLLECTION_NAME
                )
                st.write("New vector store created.")
            except Exception as e:
                st.write(f"Error creating new vector store: {e}")
                st.stop()

        # Create the QA chain using the (new or updated) vector store
        num_chunks = 4
        st.session_state.conversation = get_qa_chain(vectorestore, num_chunks)
        st.session_state.processComplete = True

    # Show the chat input field only if the QA chain is ready
    if st.session_state.processComplete:
        user_question = st.chat_input("Submit your symptoms and inquire about your medical reports.")
        if user_question:
            handel_userinput(user_question)

    # Display chat history
    if st.session_state.chat_history:
        response_container = st.container()
        with response_container:
            for i, msg in enumerate(st.session_state.chat_history):
                if i % 2 == 0:
                    message(msg, is_user=True, key=str(i))
                else:
                    message(msg, key=str(i))

# File and document processing functions

def get_files_text(uploaded_file):
    text = ""
    _, file_extension = os.path.splitext(uploaded_file.name)
    file_extension = file_extension.lower()
    if file_extension == ".pdf":
        text += get_pdf_text(uploaded_file)
    elif file_extension == ".docx":
        text += get_docx_text(uploaded_file)
    else:
        pass
    return text

def get_pdf_text(pdf):
    pdf_reader = PdfReader(pdf)
    text = ""
    for page in pdf_reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text
    return text

def get_docx_text(file):
    doc = docx.Document(file)
    allText = []
    for para in doc.paragraphs:
        allText.append(para.text)
    return " ".join(allText)

def get_text_chunks(text, filename):
    text_splitter = CharacterTextSplitter(separator="\n", chunk_size=80, chunk_overlap=20, length_function=len)
    chunks = text_splitter.split_text(text)
    doc_list = []
    for chunk in chunks:
        metadata = {"source": filename}
        doc_list.append(Document(page_content=chunk, metadata=metadata))
    return doc_list

# Create a QA chain from the vector store
def get_qa_chain(vectorstore, num_chunks):
    prompt_template = """
    You are trained to extract Answer from the given Context and Question. Then, Detail the Answer in 600 words. 
    If the Answer is not found in the Context, then return "N/A", otherwise return the detailed Answer.
    Context: {context}
    Question: {question}"""
    mprompt_url = PromptTemplate(template=prompt_template, input_variables=["context", "question"], validate_template=False)
    chain_type_kwargs = {"prompt": mprompt_url}

    qa = RetrievalQA.from_chain_type(
        llm=ChatGoogleGenerativeAI(
            model="gemini-1.5-flash", 
            temperature=1, 
            max_tokens=600, 
            api_key=genai_api_key, 
            request_options=retry_policy
        ),
        chain_type="stuff",
        retriever=vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": num_chunks}),
        chain_type_kwargs=chain_type_kwargs,
        return_source_documents=True
    )
    return qa

def handel_userinput(user_question):
    with st.spinner('Generating response...'):
        result = st.session_state.conversation({"query": user_question})
        response = result['result']
        source = result['source_documents'][0].metadata.get('source', "unknown")
    st.session_state.chat_history.append(user_question)
    st.session_state.chat_history.append(f"{response}\nSource Document: {source}")

    response_container = st.container()
    with response_container:
        for i, msg in enumerate(st.session_state.chat_history):
            if i % 2 == 0:
                message(msg, is_user=True, key=str(i))
            else:
                message(msg, key=str(i))

if __name__ == '__main__':
    main()
